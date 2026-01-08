"""Tests for base DOCX tool (markdown_to_word).

These tests verify that the markdown to Word conversion works correctly,
including headers, lists, tables, formatting, links, and block quotes.

Output files are saved to tests/output/docx/ directory for manual inspection.
"""

import sys
from pathlib import Path

# Add project root to path for imports
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

import pytest
from docx import Document

from docx_tools.helpers import (
    parse_inline_formatting,
    load_templates,
    parse_table,
    add_table_to_doc,
    process_list_items,
)
import re

# Output directory for test files
OUTPUT_DIR = Path(__file__).parent / "output" / "docx"


@pytest.fixture(scope="module", autouse=True)
def setup_output_dir():
    """Create output directory if it doesn't exist."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    yield


def create_word_document(markdown_content: str) -> Document:
    """Convert Markdown to Word document and return the Document object.

    This is a test-friendly version that returns the Document directly
    instead of saving via upload_file.
    """
    path = load_templates()

    if path:
        doc = Document(path)
    else:
        doc = Document()

    lines = markdown_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        # Check if this line ends with two spaces (line break)
        if line.endswith('  '):
            paragraph_lines = []
            while i < len(lines):
                current_line = lines[i]
                if not current_line.strip():
                    break
                paragraph_lines.append(current_line)
                i += 1
                if not current_line.endswith('  '):
                    break

            full_text = '  \n'.join(paragraph_lines)
            first_line = paragraph_lines[0].strip()

            if first_line.startswith('#'):
                header_level = len(first_line) - len(first_line.lstrip('#'))
                header_text = first_line.lstrip('#').strip()
                heading = doc.add_heading('', level=min(header_level, 6))
                parse_inline_formatting(header_text, heading)
            elif first_line.startswith('>'):
                quote_text = full_text[1:].strip()
                quote_paragraph = doc.add_paragraph()
                quote_paragraph.style = 'Quote'
                parse_inline_formatting(quote_text, quote_paragraph)
            else:
                paragraph = doc.add_paragraph()
                parse_inline_formatting(full_text, paragraph)
            continue

        line = line.strip()

        if line.startswith('#'):
            header_level = len(line) - len(line.lstrip('#'))
            header_text = line.lstrip('#').strip()
            heading = doc.add_heading('', level=min(header_level, 6))
            parse_inline_formatting(header_text, heading)
            i += 1

        elif line.startswith('|'):
            table_data, i = parse_table(lines, i)
            if table_data:
                add_table_to_doc(table_data, doc)

        elif re.match(r'^\d+\.\s+', line):
            i = process_list_items(lines, i, doc, True, 0)

        elif re.match(r'^[-*+]\s+', line):
            i = process_list_items(lines, i, doc, False, 0)

        elif line.startswith('---') or line.startswith('***'):
            doc.add_paragraph()
            i += 1

        elif line.startswith('>'):
            quote_text = line[1:].strip()
            quote_paragraph = doc.add_paragraph()
            quote_paragraph.style = 'Quote'
            parse_inline_formatting(quote_text, quote_paragraph)
            i += 1

        else:
            paragraph = doc.add_paragraph()
            parse_inline_formatting(line, paragraph)
            i += 1

    return doc


def save_test_document(markdown: str, filename: str) -> Document:
    """Convert markdown to Word and save directly to test output directory.

    Args:
        markdown: Markdown content to convert
        filename: Output filename (e.g., 'header_h1.docx')

    Returns:
        The generated Document object for assertions
    """
    doc = create_word_document(markdown)
    output_path = OUTPUT_DIR / filename
    doc.save(str(output_path))
    print(f"Saved: {output_path}")
    return doc


# =============================================================================
# Header Tests
# =============================================================================

class TestHeaders:
    """Tests for markdown headers conversion."""

    def test_h1_header(self):
        """Test H1 header conversion."""
        markdown = "# Main Title"
        doc = save_test_document(markdown, "header_h1.docx")
        assert doc is not None

    def test_h2_header(self):
        """Test H2 header conversion."""
        markdown = "## Section Title"
        doc = save_test_document(markdown, "header_h2.docx")
        assert doc is not None

    def test_h3_header(self):
        """Test H3 header conversion."""
        markdown = "### Subsection Title"
        doc = save_test_document(markdown, "header_h3.docx")
        assert doc is not None

    def test_multiple_headers(self):
        """Test document with multiple header levels."""
        markdown = """# Document Title

## Introduction

Some intro text here.

### Details

More details.

## Conclusion

Final thoughts.
"""
        doc = save_test_document(markdown, "header_multiple.docx")
        assert doc is not None

    def test_header_with_formatting(self):
        """Test header with inline formatting."""
        markdown = "# Title with **bold** and *italic*"
        doc = save_test_document(markdown, "header_formatted.docx")
        assert doc is not None


# =============================================================================
# List Tests
# =============================================================================

class TestLists:
    """Tests for markdown list conversion."""

    def test_unordered_list(self):
        """Test unordered (bullet) list."""
        markdown = """- First item
- Second item
- Third item
"""
        doc = save_test_document(markdown, "list_unordered.docx")
        assert doc is not None

    def test_ordered_list(self):
        """Test ordered (numbered) list."""
        markdown = """1. First item
2. Second item
3. Third item
"""
        doc = save_test_document(markdown, "list_ordered.docx")
        assert doc is not None

    def test_nested_list(self):
        """Test nested list items."""
        markdown = """- Main item 1
   - Sub item 1.1
   - Sub item 1.2
- Main item 2
   - Sub item 2.1
"""
        doc = save_test_document(markdown, "list_nested.docx")
        assert doc is not None

    def test_list_with_formatting(self):
        """Test list items with inline formatting."""
        markdown = """- **Bold item**
- *Italic item*
- Item with `code`
- Item with [link](https://example.com)
"""
        doc = save_test_document(markdown, "list_formatted.docx")
        assert doc is not None

    def test_mixed_list_types(self):
        """Test document with both ordered and unordered lists."""
        markdown = """## Shopping List

- Apples
- Bananas
- Oranges

## Steps to Follow

1. First step
2. Second step
3. Third step
"""
        doc = save_test_document(markdown, "list_mixed.docx")
        assert doc is not None


# =============================================================================
# Table Tests
# =============================================================================

class TestTables:
    """Tests for markdown table conversion."""

    def test_simple_table(self):
        """Test simple table conversion."""
        markdown = """| Name | Age | City |
|------|-----|------|
| John | 25  | NYC  |
| Jane | 30  | LA   |
"""
        doc = save_test_document(markdown, "table_simple.docx")
        assert doc is not None

    def test_table_with_formatting(self):
        """Test table with formatted cells."""
        markdown = """| Feature | Description |
|---------|-------------|
| **Bold** | This is bold |
| *Italic* | This is italic |
| `Code` | This is code |
"""
        doc = save_test_document(markdown, "table_formatted.docx")
        assert doc is not None

    def test_table_with_alignment(self):
        """Test table with column alignment markers."""
        markdown = """| Left | Center | Right |
|:-----|:------:|------:|
| L1   | C1     | R1    |
| L2   | C2     | R2    |
"""
        doc = save_test_document(markdown, "table_aligned.docx")
        assert doc is not None


# =============================================================================
# Inline Formatting Tests
# =============================================================================

class TestInlineFormatting:
    """Tests for inline markdown formatting."""

    def test_bold_text(self):
        """Test bold text conversion."""
        markdown = "This is **bold** text."
        doc = save_test_document(markdown, "format_bold.docx")
        assert doc is not None

    def test_italic_text(self):
        """Test italic text conversion."""
        markdown = "This is *italic* text."
        doc = save_test_document(markdown, "format_italic.docx")
        assert doc is not None

    def test_inline_code(self):
        """Test inline code conversion."""
        markdown = "Use the `print()` function."
        doc = save_test_document(markdown, "format_code.docx")
        assert doc is not None

    def test_hyperlink(self):
        """Test hyperlink conversion."""
        markdown = "Visit [our website](https://example.com) for more info."
        doc = save_test_document(markdown, "format_link.docx")
        assert doc is not None

    def test_mixed_formatting(self):
        """Test multiple formatting types in one paragraph."""
        markdown = "This has **bold**, *italic*, `code`, and [link](https://test.com)."
        doc = save_test_document(markdown, "format_mixed.docx")
        assert doc is not None

    def test_nested_formatting(self):
        """Test nested formatting (bold containing italic)."""
        markdown = "This is **bold with *italic* inside**."
        doc = save_test_document(markdown, "format_nested.docx")
        assert doc is not None

    def test_escaped_characters(self):
        """Test escaped markdown characters."""
        markdown = r"This has \*asterisks\* and \**double asterisks\**."
        doc = save_test_document(markdown, "format_escaped.docx")
        assert doc is not None


# =============================================================================
# Block Quote Tests
# =============================================================================

class TestBlockQuotes:
    """Tests for block quote conversion."""

    def test_simple_quote(self):
        """Test simple block quote."""
        markdown = "> This is a quoted text."
        doc = save_test_document(markdown, "quote_simple.docx")
        assert doc is not None

    def test_quote_with_formatting(self):
        """Test block quote with inline formatting."""
        markdown = "> This quote has **bold** and *italic* text."
        doc = save_test_document(markdown, "quote_formatted.docx")
        assert doc is not None


# =============================================================================
# Complex Document Tests
# =============================================================================

class TestComplexDocuments:
    """Tests for complex documents combining multiple elements."""

    def test_full_document(self):
        """Test a complete document with all elements."""
        markdown = """# Project Report

## Executive Summary

This report provides a **comprehensive analysis** of the project status.

## Key Findings

The following points summarize our findings:

- Revenue increased by **15%**
- Customer satisfaction improved to *92%*
- New features deployed successfully

## Data Overview

| Metric | Q1 | Q2 | Q3 |
|--------|----|----|-----|
| Sales | 100 | 120 | 150 |
| Users | 500 | 600 | 800 |

## Next Steps

1. Expand into new markets
2. Invest in R&D
3. Focus on customer retention

> "The best way to predict the future is to create it." - Peter Drucker

## Conclusion

Visit [our dashboard](https://example.com/dashboard) for live updates.
"""
        doc = save_test_document(markdown, "complex_full_document.docx")
        assert doc is not None

    def test_legal_contract_style(self):
        """Test legal contract style document with numbered sections."""
        markdown = """# SERVICE AGREEMENT

1. PARTIES
   - This agreement is between Company A and Company B.
   - Both parties agree to the following terms.

2. SERVICES
   - Company A will provide consulting services.
   - Services include analysis, recommendations, and implementation support.

3. PAYMENT TERMS
   - Payment is due within 30 days of invoice.
   - Late payments incur a 1.5% monthly fee.

4. CONFIDENTIALITY
   - Both parties agree to maintain confidentiality.
   - This obligation survives termination of the agreement.
"""
        doc = save_test_document(markdown, "complex_contract.docx")
        assert doc is not None

    def test_technical_documentation(self):
        """Test technical documentation style."""
        markdown = """# API Documentation

## Authentication

All API requests require authentication using an API key.

Use the `Authorization` header:

> Authorization: Bearer YOUR_API_KEY

## Endpoints

### GET /users

Returns a list of users.

| Parameter | Type | Required | Description |
|-----------|------|----------|-------------|
| page | integer | No | Page number |
| limit | integer | No | Items per page |

### POST /users

Creates a new user.

**Request Body:**

- `name` - User's full name
- `email` - User's email address
- `role` - User's role (*admin*, *user*, or *guest*)
"""
        doc = save_test_document(markdown, "complex_api_docs.docx")
        assert doc is not None


# =============================================================================
# Edge Cases
# =============================================================================

class TestEdgeCases:
    """Tests for edge cases and special scenarios."""

    def test_empty_content(self):
        """Test with empty content."""
        markdown = ""
        doc = save_test_document(markdown, "edge_empty_content.docx")
        assert doc is not None

    def test_only_whitespace(self):
        """Test with only whitespace."""
        markdown = "   \n\n   \n"
        doc = save_test_document(markdown, "edge_only_whitespace.docx")
        assert doc is not None

    def test_multiple_empty_lines(self):
        """Test preservation of multiple empty lines."""
        markdown = """First paragraph.


Third paragraph (after two empty lines).
"""
        doc = save_test_document(markdown, "edge_empty_lines.docx")
        assert doc is not None

    def test_unicode_content(self):
        """Test with unicode characters."""
        markdown = """# Vícejazyčný dokument

Příliš žluťoučký kůň úpěl ďábelské ódy.

日本語テキスト

Emoji: 👋 🌍 ✨
"""
        doc = save_test_document(markdown, "edge_unicode.docx")
        assert doc is not None

    def test_long_paragraph(self):
        """Test with very long paragraph."""
        long_text = "Lorem ipsum dolor sit amet. " * 50
        markdown = f"# Long Document\n\n{long_text}"
        doc = save_test_document(markdown, "edge_long_paragraph.docx")
        assert doc is not None

    def test_special_xml_characters(self):
        """Test with characters that need XML escaping."""
        markdown = "This has < and > and & characters."
        doc = save_test_document(markdown, "edge_xml_chars.docx")
        assert doc is not None

    def test_line_breaks(self):
        """Test soft line breaks (two spaces at end)."""
        markdown = """This is line one.  
This is line two (same paragraph).  
This is line three.
"""
        doc = save_test_document(markdown, "edge_line_breaks.docx")
        assert doc is not None


# =============================================================================
# Regression Tests for helpers.py changes
# =============================================================================

class TestHelpersRegression:
    """Regression tests for helpers.py functionality used by base tool."""

    def test_parse_inline_formatting_plain(self):
        """Test parse_inline_formatting with plain text."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("Plain text", para)
        assert para.text == "Plain text"

    def test_parse_inline_formatting_bold(self):
        """Test parse_inline_formatting with bold."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("Text with **bold** word", para)
        assert "bold" in para.text
        bold_runs = [r for r in para.runs if r.bold]
        assert len(bold_runs) > 0

    def test_parse_inline_formatting_italic(self):
        """Test parse_inline_formatting with italic."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("Text with *italic* word", para)
        assert "italic" in para.text
        italic_runs = [r for r in para.runs if r.italic]
        assert len(italic_runs) > 0

    def test_parse_inline_formatting_code(self):
        """Test parse_inline_formatting with inline code."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("Use `code` here", para)
        assert "code" in para.text
        code_runs = [r for r in para.runs if r.font.name == "Courier New"]
        assert len(code_runs) > 0

    def test_parse_inline_formatting_link(self):
        """Test parse_inline_formatting with hyperlink."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("Visit [link](https://example.com)", para)
        # Check that hyperlink element exists
        hyperlinks = para._p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
        assert len(hyperlinks) > 0

    def test_parse_inline_formatting_nested(self):
        """Test parse_inline_formatting with nested formatting."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("This is **bold with *italic* inside**", para)

        # Should have runs with both bold and italic
        bold_italic_runs = [r for r in para.runs if r.bold and r.italic]
        assert len(bold_italic_runs) > 0

    def test_parse_inline_formatting_multiple_bold(self):
        """Test parse_inline_formatting with multiple bold sections."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("**First** and **second** bold", para)

        bold_runs = [r for r in para.runs if r.bold and r.text.strip()]
        assert len(bold_runs) >= 2


# =============================================================================
# Comprehensive Visual Test
# =============================================================================

class TestVisualInspection:
    """Comprehensive test for manual visual inspection of generated documents.

    This test creates a single document with ALL supported markdown features
    for easy visual verification in Word/LibreOffice.

    Output: tests/output/docx/base/VISUAL_INSPECTION_comprehensive.docx
    """

    def test_comprehensive_visual_document(self):
        """Generate a comprehensive document for visual inspection.

        This document includes:
        - All heading levels (H1-H6)
        - Paragraphs with various inline formatting
        - Ordered and unordered lists (including nested)
        - Tables with formatting
        - Block quotes
        - Hyperlinks
        - Unicode and special characters
        - Line breaks
        """
        markdown = """# Comprehensive Visual Inspection Document

This document is designed for **manual visual inspection** to verify that all markdown 
features are correctly converted to Word format. Open this file in Microsoft Word or 
LibreOffice Writer to check the formatting.

---

## 1. Heading Levels

### Heading Level 3

#### Heading Level 4

##### Heading Level 5

###### Heading Level 6

---

## 2. Inline Formatting

This paragraph contains **bold text**, *italic text*, and ***bold italic text***. 
You can also use `inline code` for technical terms like `print()` or `variable_name`.

Here is a [hyperlink to example.com](https://example.com) and another 
[link to Google](https://www.google.com).

Mixed formatting: **bold with *nested italic* inside** and *italic with **nested bold** inside*.

---

## 3. Unordered Lists

Simple bullet list:

- First item
- Second item with **bold** text
- Third item with *italic* text
- Fourth item with `code`
- Fifth item with [link](https://example.com)

Nested bullet list:

- Main item 1
   - Sub-item 1.1
   - Sub-item 1.2
      - Deep nested item
   - Sub-item 1.3
- Main item 2
   - Sub-item 2.1

Different markers (should all render as bullets):

* Asterisk item 1
* Asterisk item 2

+ Plus item 1
+ Plus item 2

---

## 4. Ordered Lists

Simple numbered list:

1. First step
2. Second step with **important** info
3. Third step with *emphasis*
4. Fourth step with `code snippet`

Nested numbered list:

1. Main step 1
   1. Sub-step 1.1
   2. Sub-step 1.2
2. Main step 2
   1. Sub-step 2.1
   2. Sub-step 2.2
   3. Sub-step 2.3
3. Main step 3

---

## 5. Mixed List Types

Shopping list:

- Apples
- Bananas
- Oranges

Preparation steps:

1. Wash the fruit
2. Cut into pieces
3. Serve and enjoy

---

## 6. Tables

### Simple Table

| Name | Age | City |
|------|-----|------|
| John | 25 | New York |
| Jane | 30 | Los Angeles |
| Bob | 35 | Chicago |

### Table with Formatting

| Feature | Description | Status |
|---------|-------------|--------|
| **Bold Feature** | This feature is *very important* | Active |
| *Italic Feature* | Contains `code` elements | Pending |
| Regular Feature | Visit [docs](https://docs.example.com) | Complete |

### Table with Alignment

| Left Aligned | Center Aligned | Right Aligned |
|:-------------|:--------------:|--------------:|
| L1 | C1 | R1 |
| L2 | C2 | R2 |
| L3 | C3 | R3 |

---

## 7. Block Quotes

> This is a simple block quote.

> This block quote contains **bold** and *italic* formatting.

> "The best way to predict the future is to create it." - Peter Drucker

---

## 8. Unicode and Special Characters

### Czech Text
Příliš žluťoučký kůň úpěl ďábelské ódy.

### German Text
Größe, Müller, Straße, Übung

### Japanese Text
こんにちは世界 (Hello World)

### Emoji
Hello 👋 World 🌍 Stars ⭐✨ Check ✓ Heart ❤️

### Special XML Characters
5 > 3 and 2 < 4 and A & B

---

## 9. Line Breaks

This is line one.  
This is line two (same paragraph, soft break).  
This is line three (still same paragraph).

---

## 10. Complex Paragraph

This paragraph demonstrates **multiple formatting options** combined together. 
We have *italic text*, `inline code`, and [hyperlinks](https://example.com). 
You can even have **bold with *nested italic*** or *italic with **nested bold***. 
Special characters like < > & are properly escaped.

---

## 11. Technical Documentation Style

### API Endpoint: GET /users

Returns a list of users.

**Parameters:**

| Parameter | Type | Required | Description |
|-----------|------|----------|-------------|
| `page` | integer | No | Page number (default: 1) |
| `limit` | integer | No | Items per page (default: 20) |
| `sort` | string | No | Sort field |

**Example Response:**

> The response includes user data in JSON format.

---

## 12. Legal Document Style

1. PARTIES
   - This agreement is between **Company A** and **Company B**.
   - Both parties agree to the terms below.

2. TERMS AND CONDITIONS
   - All payments due within *30 days*.
   - Late payments incur a `1.5%` monthly fee.

3. CONFIDENTIALITY
   - Both parties maintain strict confidentiality.
   - See [Privacy Policy](https://example.com/privacy) for details.

---

## Conclusion

This document contains all supported markdown elements. If you can read this 
and all formatting above appears correct, the markdown-to-Word conversion is 
working properly! 🎉

**Document generated for visual inspection purposes.**

*Last updated: January 2026*
"""
        doc = save_test_document(markdown, "VISUAL_INSPECTION_comprehensive.docx")
        assert doc is not None

        # Basic sanity checks
        assert len(doc.paragraphs) > 50, "Document should have many paragraphs"

        # Check for various elements in paragraphs
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Comprehensive Visual Inspection" in full_text
        assert "bold text" in full_text
        assert "italic text" in full_text
        assert "First item" in full_text
        assert "žluťoučký" in full_text  # Czech unicode
        assert "こんにちは" in full_text  # Japanese

        # Check tables exist and have content
        assert len(doc.tables) >= 3, "Document should have at least 3 tables"
        # Check table content
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "
        assert "John" in table_text  # From simple table


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])

