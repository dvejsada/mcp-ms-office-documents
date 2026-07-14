"""Integration tests for the ``procesni_podani`` dynamic DOCX template.

The template (``custom_templates/procesni_podani_template.docx`` +
``config/docx_templates.yaml``) produces a Czech court filing: head block with
court and parties, conditional přílohy/důkazy lists, a markdown body mapped
onto the firm's section/paragraph numbering styles, and a conditional petit.

These tests exercise the same pipeline the registered MCP tool runs
(conditional resolution + placeholder substitution with the template's style
map) and assert on styles and numbering XML, since Word computes the visible
numbers at display time.
"""
import sys
from pathlib import Path

import pytest
import yaml
from docx import Document

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx_tools.conditionals import resolve_conditionals  # noqa: E402
from docx_tools.dynamic_docx_tools import (  # noqa: E402
    find_docx_template_by_name,
    replace_placeholders_in_document,
)
from docx_tools.style_map import build_style_map  # noqa: E402

OUTPUT_DIR = Path(__file__).parent / "output" / "docx"

# Styles the template's style_mapping and argument descriptions rely on.
REQUIRED_STYLES = [
    "Oddíl_číslování", "Pododdíl_číslování", "Odstavec_číslování",
    "Pododstavec", "Petit", "Důkazy", "Důkazy odrážky", "Seznam přílohy",
    "Seznam důkazy", "PODPIS", "Quote", "List Number", "List Bullet",
    "Table Grid",
]


@pytest.fixture(scope="module")
def spec():
    cfg = yaml.safe_load(
        (project_root / "config" / "docx_templates.yaml").read_text(encoding="utf-8"))
    matches = [t for t in cfg["templates"] if t.get("name") == "procesni_podani"]
    assert len(matches) == 1, "procesni_podani template must be defined exactly once"
    return matches[0]


@pytest.fixture(scope="module")
def template_path(spec):
    path = find_docx_template_by_name(spec["docx_path"])
    assert path, f"template file {spec['docx_path']} not found"
    return path


def _payload(spec, **overrides):
    payload = {a["name"]: a.get("default") for a in spec["args"]}
    payload.update({
        "soud": "Městskému soudu v Praze<br>Slezská 2000/9<br>120 00 Praha 2",
        "misto_datum": "Praha 26. května 2026",
        "zalobce": "ABC s.r.o., IČO: 123 45 678<br>se sídlem Dlouhá 1, Praha 1",
        "zalovany": "XYZ a.s., IČO: 876 54 321<br>se sídlem Krátká 2, Brno",
        "typ_rizeni": "o zaplacení 1 000 000 Kč s příslušenstvím",
        "nazev_podani": "VYJÁDŘENÍ ŽALOVANÉHO K ŽALOBĚ",
        "podpis": "ABC advokáti s.r.o.<br>JUDr. Jan Novák, advokát",
        "text_podani": (
            "# SKUTKOVÝ STAV\n\n"
            "1. První odstavec podání.\n\n"
            "2. Druhý odstavec podání.\n\n"
            "<!-- style: Důkazy -->\nDůkaz:\n\n"
            "<!-- style: Důkazy odrážky -->\n"
            "- smlouva ze dne 1. 1. 2026\n"
            "- výslech svědka Jana Nováka\n\n"
            "> Citace smluvního ujednání.\n\n"
            "# PRÁVNÍ POSOUZENÍ\n\n"
            "3. Třetí odstavec podání.\n\n"
            "# NÁVRH VÝROKU ROZHODNUTÍ\n\n"
            "4. **S ohledem na výše uvedené žalobce navrhuje, aby nadepsaný "
            "soud vydal tento**"
        ),
        "petit": (
            "<!-- style: Petit -->\nŽalovaný je povinen zaplatit žalobci "
            "částku 1 000 000 Kč do 3 dnů od právní moci rozsudku.\n\n"
            "<!-- style: Petit -->\nŽalovaný je povinen nahradit žalobci "
            "náklady řízení."
        ),
        "prilohy": ("<!-- style: Seznam přílohy -->\n1. Plná moc\n"
                    "2. Smlouva ze dne 1. 1. 2026"),
    })
    payload.update(overrides)
    return payload


def _render(template_path, spec, payload):
    doc = Document(template_path)
    resolve_conditionals(doc, payload)
    context = {k: ("" if v is None else str(v)) for k, v in payload.items()}
    replace_placeholders_in_document(
        doc, context, build_style_map(spec.get("style_mapping")))
    return doc


def _all_texts(doc):
    texts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            texts.extend(p.text for p in ncell.paragraphs)
    for section in doc.sections:
        for part in (section.footer, section.first_page_footer):
            if part is None:
                continue
            texts.extend(p.text for p in part.paragraphs)
            for tbl in part.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        texts.extend(p.text for p in cell.paragraphs)
    return texts


def _styled(doc, style_id):
    return [p for p in doc.paragraphs if p.style is not None
            and p.style.style_id == style_id]


def test_template_defines_required_styles(template_path):
    doc = Document(template_path)
    names = {s.name for s in doc.styles}
    missing = [n for n in REQUIRED_STYLES if n not in names]
    assert not missing, f"template is missing styles: {missing}"


def test_full_filing_renders_without_leftover_placeholders(template_path, spec):
    payload = _payload(spec, ma_dukazy=True,
                       dukazy="<!-- style: Seznam důkazy -->\n1. smlouva")
    doc = _render(template_path, spec, payload)
    leftovers = [t for t in _all_texts(doc) if "{{" in t or "}}" in t]
    assert not leftovers, f"unreplaced placeholders: {leftovers}"

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DIR / "procesni_podani_full.docx"))


def test_body_maps_markdown_onto_firm_styles(template_path, spec):
    doc = _render(template_path, spec, _payload(spec))

    headings = _styled(doc, "Oddlslovn")
    assert [h.text for h in headings] == [
        "SKUTKOVÝ STAV", "PRÁVNÍ POSOUZENÍ", "NÁVRH VÝROKU ROZHODNUTÍ"]

    items = _styled(doc, "Odstavecslovn")
    assert len(items) == 4
    # Numbered paragraphs continue across sections and across the interposed
    # Důkaz note and citation: the later list instances resume via
    # startOverride instead of restarting at 1.
    num_root = doc.part.numbering_part.element
    overrides = []
    for p in items:
        num_ids = p._p.xpath(".//w:numPr/w:numId/@w:val")
        assert num_ids, f"numbered paragraph lacks numbering: {p.text!r}"
        num = num_root.num_having_numId(int(num_ids[0]))
        vals = num.xpath('./w:lvlOverride[@w:ilvl="0"]/w:startOverride/@w:val')
        overrides.append(vals[0] if vals else None)
    assert overrides == ["1", "1", "3", "4"]

    # Evidence: a "Důkaz:" label followed by a bulleted list of the items.
    assert [p.text for p in _styled(doc, "Dkazy")] == ["Důkaz:"]
    bullets = _styled(doc, "Dkazyodrazky")
    assert [p.text for p in bullets] == [
        "smlouva ze dne 1. 1. 2026", "výslech svědka Jana Nováka"]
    assert [p.text for p in _styled(doc, "Quote")] == ["Citace smluvního ujednání."]


def test_petit_and_prilohy_use_directive_styles(template_path, spec):
    doc = _render(template_path, spec, _payload(spec))

    petit = _styled(doc, "Petit")
    assert len(petit) == 2
    # Petit relies on its style's numbering (upper-Roman I., II.), so the
    # directive-styled paragraphs must not carry an explicit numPr override.
    for p in petit:
        assert not p._p.xpath(".//w:numPr/w:numId"), \
            "explicit numbering would shadow the Petit style's Roman numerals"

    prilohy = _styled(doc, "Seznamplohy")
    assert [p.text for p in prilohy] == ["Plná moc", "Smlouva ze dne 1. 1. 2026"]


def _page_breaks(doc):
    return len(doc.element.body.xpath('.//w:br[@w:type="page"]'))


def test_conditional_blocks(template_path, spec):
    # Default flags: přílohy in, důkazy out, petit in.
    doc = _render(template_path, spec, _payload(spec))
    texts = _all_texts(doc)
    assert "Přílohy:" in texts
    assert "Bez příloh." not in texts
    assert "Seznam důkazů:" not in texts
    # Without the evidence list there is no explicit page break; the body
    # still starts on a fresh page via the cover section's break.
    assert _page_breaks(doc) == 0
    assert len(doc.element.body.xpath('.//w:pPr/w:sectPr')) == 1

    # Flip the flags: no attachments, no petit, with an evidence list.
    doc = _render(template_path, spec, _payload(
        spec, ma_prilohy=False, ma_petit=False, ma_dukazy=True,
        dukazy="<!-- style: Seznam důkazy -->\n1. smlouva", petit=" "))
    texts = _all_texts(doc)
    assert "Bez příloh." in texts
    assert "Přílohy:" not in texts
    assert "Seznam důkazů:" in texts
    assert not _styled(doc, "Petit")
    assert "R O Z S U D E K" not in texts
    # The evidence list takes its own page (explicit break before it).
    assert _page_breaks(doc) == 1
