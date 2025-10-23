import io
import logging
import re
from docx import Document

from upload_tools import upload_file
from .helpers import (
    load_templates,
    parse_inline_formatting,
    parse_table,
    add_table_to_doc,
    process_list_items,
)

logger = logging.getLogger(__name__)


def markdown_to_word(markdown_content):
    """Convert Markdown to Word document."""
    logger.info("Starting markdown_to_word conversion")
    path = load_templates()

    # Create document with or without template
    if path:
        logger.debug(f"Using Word template at: {path}")
        doc = Document(path)
    else:
        doc = Document()  # Create blank document if no template
        logger.warning("No template found, creating blank document")

    # Split content into lines, but preserve line breaks within paragraphs
    lines = markdown_content.split('\n')
    i = 0

    # Simple parsing counters for summary
    headers_count = 0
    tables_count = 0
    ordered_lists = 0
    unordered_lists = 0
    quotes_count = 0
    paragraphs_count = 0

    try:
        while i < len(lines):
            line = lines[i]

            # Handle multiple empty lines (preserve spacing)
            if not line.strip():
                empty_line_count = 0
                start_empty = i

                # Count consecutive empty lines
                while i < len(lines) and not lines[i].strip():
                    empty_line_count += 1
                    i += 1

                # Add appropriate spacing based on number of empty lines
                if empty_line_count == 1:
                    pass
                elif empty_line_count >= 2:
                    for _ in range(empty_line_count - 1):
                        doc.add_paragraph()
                        paragraphs_count += 1
                continue

            # Check if this line ends with two spaces (line break)
            if line.endswith('  '):
                # Collect lines that are part of the same paragraph (connected by line breaks)
                paragraph_lines = []
                while i < len(lines):
                    current_line = lines[i]
                    if not current_line.strip():
                        break

                    paragraph_lines.append(current_line)
                    i += 1

                    if not current_line.endswith('  '):
                        break

                full_text = '  \n'.join(paragraph_lines)
                first_line = paragraph_lines[0].strip()

                if first_line.startswith('#'):
                    header_level = len(first_line) - len(first_line.lstrip('#'))
                    header_text = first_line.lstrip('#').strip()
                    heading = doc.add_heading('', level=min(header_level, 6))
                    parse_inline_formatting(header_text, heading)
                    headers_count += 1
                    logger.debug(f"Header (level {header_level}): {header_text}")
                elif first_line.startswith('>'):
                    quote_text = full_text[1:].strip()
                    quote_paragraph = doc.add_paragraph()
                    quote_paragraph.style = 'Quote'
                    parse_inline_formatting(quote_text, quote_paragraph)
                    quotes_count += 1
                else:
                    paragraph = doc.add_paragraph()
                    parse_inline_formatting(full_text, paragraph)
                    paragraphs_count += 1
                continue

            line = line.strip()

            if line.startswith('#'):
                header_level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('#').strip()
                heading = doc.add_heading('', level=min(header_level, 6))
                parse_inline_formatting(header_text, heading)
                headers_count += 1
                logger.debug(f"Header (level {header_level}): {header_text}")
                i += 1

            elif line.startswith('|'):
                table_data, i = parse_table(lines, i)
                if table_data:
                    add_table_to_doc(table_data, doc)
                    tables_count += 1
                    logger.debug(f"Added table with {len(table_data)} rows")

            elif re.match(r'^\d+\.\s+', line):
                i = process_list_items(lines, i, doc, True, 0)
                ordered_lists += 1

            elif re.match(r'^[-*+]\s+', line):
                i = process_list_items(lines, i, doc, False, 0)
                unordered_lists += 1

            elif line.startswith('---') or line.startswith('***'):
                doc.add_paragraph()
                paragraphs_count += 1
                i += 1

            elif line.startswith('>'):
                quote_text = line[1:].strip()
                quote_paragraph = doc.add_paragraph()
                quote_paragraph.style = 'Quote'
                parse_inline_formatting(quote_text, quote_paragraph)
                quotes_count += 1
                i += 1

            else:
                paragraph = doc.add_paragraph()
                parse_inline_formatting(line, paragraph)
                paragraphs_count += 1
                i += 1

    except Exception as e:
        logger.error(f"Error in parsing markdown: {e}", exc_info=True)
        return f"Error in parsing markdown: {e}"

    # Save the document to BytesIO and upload
    try:
        logger.info("Saving Word document to memory buffer")
        file_object = io.BytesIO()
        doc.save(file_object)
        file_object.seek(0)

        result = upload_file(file_object, "docx")
        file_object.close()

        logger.info(
            f"Word upload completed (headers={headers_count}, tables={tables_count}, ordered_lists={ordered_lists}, "
            f"unordered_lists={unordered_lists}, quotes={quotes_count}, paragraphs={paragraphs_count})"
        )
        return result
    except Exception as e:
        logger.error(f"Error saving/uploading Word document: {e}", exc_info=True)
        return f"Error saving/uploading Word document: {e}"
