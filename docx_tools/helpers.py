import logging
import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from template_utils import find_docx_template

logger = logging.getLogger(__name__)


def load_templates():
    """Resolve Word template path from custom/default template directories.

    Returns absolute path as string or None if not found.
    """
    path = find_docx_template()
    if path:
        logger.debug(f"Using Word template: {path}")
    else:
        logger.warning("No Word template found, will create a blank document")
    return path


def add_hyperlink(paragraph, text, url, color="0000FF", underline=True):
    """Adds a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    new_run.append(rPr)

    # Create the text element properly
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    # Preserve spaces at start/end
    text_elem.set(qn('xml:space'), 'preserve')
    new_run.append(text_elem)

    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def parse_inline_formatting(text, paragraph, bold=False, italic=False):
    """Parse inline markdown formatting like **bold**, *italic*, and [links](url)

    Args:
        text: The text to parse
        paragraph: The paragraph to add runs to
        bold: Whether the current context is bold (for nested formatting)
        italic: Whether the current context is italic (for nested formatting)
    """
    # First handle escape characters
    text = handle_escapes(text)

    # Handle line breaks (two spaces at end of line)
    # Split by line breaks while preserving them
    line_parts = text.split('  \n')

    for line_idx, line_part in enumerate(line_parts):
        if not line_part and line_idx == len(line_parts) - 1:
            continue

        # Split text by formatting markers while preserving the markers
        # Regex explanation:
        # - \*\*(?:[^*]|\*(?!\*))+\*\* : bold (**...**) - matches ** followed by any chars except **, ending with **
        # - \*(?:[^*]|\*\*[^*]*\*\*)+\* : italic (*...*) - matches * followed by any chars or nested **, ending with *
        # - `[^`]+` : inline code
        # - \[[^\]]*\]\([^)]*\) : links [text](url)
        parts = re.split(r'(\*\*(?:[^*]|\*(?!\*))+\*\*|\*(?:[^*]|\*\*[^*]*\*\*)+\*|`[^`]+`|\[[^\]]*\]\([^)]*\))', line_part)

        for part in parts:
            if not part:
                continue

            # Bold text (**text**)
            if part.startswith('**') and part.endswith('**'):
                inner_text = part[2:-2]
                # Recursively parse inner content for nested formatting (e.g., *italic* inside bold)
                _parse_with_formatting(inner_text, paragraph, bold=True, italic=italic)
            # Italic text (*text*)
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                inner_text = part[1:-1]
                # Recursively parse inner content for nested formatting (e.g., **bold** inside italic)
                _parse_with_formatting(inner_text, paragraph, bold=bold, italic=True)
            # Inline code (`code`)
            elif part.startswith('`') and part.endswith('`'):
                code_text = part[1:-1]
                run = paragraph.add_run(code_text)
                run.font.name = 'Courier New'
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
            # Links [text](url)
            elif part.startswith('[') and '](' in part and part.endswith(')'):
                link_match = re.match(r'\[(.*?)]\((.*?)\)', part)
                if link_match:
                    link_text, url = link_match.groups()
                    add_hyperlink(paragraph, link_text, url)
            else:
                # Plain text - apply current formatting context
                run = paragraph.add_run(part)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True

        # Add line break if this isn't the last part
        if line_idx < len(line_parts) - 1:
            paragraph.add_run().add_break()


def _parse_with_formatting(text, paragraph, bold=False, italic=False):
    """Parse text that may contain nested formatting markers.

    This is a helper for recursive parsing of nested markdown.
    """
    # Split by formatting markers (same regex as in parse_inline_formatting)
    parts = re.split(r'(\*\*(?:[^*]|\*(?!\*))+\*\*|\*(?:[^*]|\*\*[^*]*\*\*)+\*|`[^`]+`|\[[^\]]*\]\([^)]*\))', text)

    for part in parts:
        if not part:
            continue

        # Bold text (**text**) - nested inside current context
        if part.startswith('**') and part.endswith('**'):
            inner_text = part[2:-2]
            # Further nesting - parse inner content
            _parse_with_formatting(inner_text, paragraph, bold=True, italic=italic)
        # Italic text (*text*) - nested inside current context
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            inner_text = part[1:-1]
            _parse_with_formatting(inner_text, paragraph, bold=bold, italic=True)
        # Inline code (`code`)
        elif part.startswith('`') and part.endswith('`'):
            code_text = part[1:-1]
            run = paragraph.add_run(code_text)
            run.font.name = 'Courier New'
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        # Links [text](url)
        elif part.startswith('[') and '](' in part and part.endswith(')'):
            link_match = re.match(r'\[(.*?)]\((.*?)\)', part)
            if link_match:
                link_text, url = link_match.groups()
                add_hyperlink(paragraph, link_text, url)
        else:
            # Plain text with inherited formatting
            run = paragraph.add_run(part)
            if bold:
                run.bold = True
            if italic:
                run.italic = True


def handle_escapes(text):
    """Handle backslash escaped characters"""
    # Replace escaped characters with placeholders first to avoid conflicts
    escape_map = {}
    placeholder_counter = 0

    def replace_escape(match):
        nonlocal placeholder_counter
        escaped_char = match.group(1)
        placeholder = f"__ESC_{placeholder_counter}__"
        escape_map[placeholder] = escaped_char
        placeholder_counter += 1
        return placeholder

    # Find and replace all escaped characters
    text = re.sub(r'\\(.)', replace_escape, text)

    # After all other processing, restore the escaped characters
    for placeholder, char in escape_map.items():
        text = text.replace(placeholder, char)

    return text


def parse_table(lines, start_idx):
    """Parse markdown table and return the table data and next line index"""
    table_lines = []
    i = start_idx

    # Find all table lines
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            table_lines.append(line)
            i += 1
        else:
            break

    if len(table_lines) < 2:  # Need at least header and separator
        return None, start_idx + 1

    # Parse table data
    table_data = []
    for line in table_lines:
        # Skip separator line (contains dashes)
        if '---' in line or ':-:' in line or ':--' in line or '--:' in line:
            continue

        # Split by | and clean up
        cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last
        table_data.append(cells)

    return table_data, i


def add_table_to_doc(table_data, doc):
    """Add table data to Word document"""
    if not table_data:
        return

    rows = len(table_data)
    cols = max(len(row) for row in table_data) if table_data else 0

    word_table = doc.add_table(rows=rows, cols=cols)
    word_table.style = 'Table Grid'

    for i, row_data in enumerate(table_data):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                cell = word_table.cell(i, j)
                if cell.paragraphs:
                    cell.paragraphs[0].clear()

                cell_paragraph = cell.paragraphs[0]
                parse_inline_formatting(cell_text, cell_paragraph)


def process_list_items(lines, start_idx, doc, is_ordered=False, level=0):
    """Process markdown list items with proper Word numbering.

    This function directly adds paragraphs to the document.

    Args:
        lines: All lines of markdown content
        start_idx: Starting index in lines
        doc: The Word document
        is_ordered: Whether this is an ordered (numbered) list
        level: Current nesting level

    Returns:
        Next line index after processing the list
    """
    i, _ = process_list_items_returning_elements(
        lines, start_idx, doc, is_ordered, level, return_elements=False
    )
    return i


def process_list_items_returning_elements(
    lines, start_idx, doc, is_ordered=False, level=0, return_elements=True
):
    """Process markdown list items and optionally return paragraph elements.

    This is the core list processing function that can either:
    - Add paragraphs directly to document (return_elements=False)
    - Return paragraph elements for manual insertion (return_elements=True)

    Args:
        lines: All lines of markdown content
        start_idx: Starting index in lines
        doc: The Word document
        is_ordered: Whether this is an ordered (numbered) list
        level: Current nesting level
        return_elements: If True, remove elements from doc and return them

    Returns:
        Tuple of (next_index, list_of_paragraph_elements) if return_elements=True
        Tuple of (next_index, None) if return_elements=False
    """
    bullet_styles = ['List Bullet', 'List Bullet 2', 'List Bullet 3']
    number_styles = ['List Number', 'List Number 2', 'List Number 3']

    style_array = number_styles if is_ordered else bullet_styles
    style = style_array[min(level, len(style_array) - 1)]

    elements = [] if return_elements else None
    i = start_idx

    while i < len(lines):
        line = lines[i].strip()

        # Determine indentation level from original line
        original_line = lines[i]
        indent = len(original_line) - len(original_line.lstrip())
        current_level = indent // 3  # Use 3 spaces per level to match typical markdown indentation

        # If indentation doesn't match our expected level, this item doesn't belong to this list
        if current_level != level:
            break

        # Check if this is a list item at our current level
        if is_ordered:
            list_match = re.match(r'^\d+\.\s+(.+)', line)
        else:
            list_match = re.match(r'^[-*+]\s+(.+)', line)

        if not list_match:
            break

        item_text = list_match.group(1)

        # Use Word's built-in list formatting - it handles numbering restart automatically
        paragraph = doc.add_paragraph(style=style)
        parse_inline_formatting(item_text, paragraph)

        if return_elements:
            elements.append(paragraph._p)
            doc._body._body.remove(paragraph._p)

        i += 1

        # Look ahead for nested items
        while i < len(lines):
            if i >= len(lines):
                break

            next_line = lines[i].strip()
            if not next_line:
                i += 1
                continue

            next_original = lines[i]
            next_indent = len(next_original) - len(next_original.lstrip())
            next_level = next_indent // 3  # Use 3 spaces per level

            if next_level > level:
                # This is a nested item - process the nested list
                if re.match(r'^\d+\.\s+', next_line):
                    i, nested_elements = process_list_items_returning_elements(
                        lines, i, doc, True, next_level, return_elements
                    )
                    if return_elements and nested_elements:
                        elements.extend(nested_elements)
                elif re.match(r'^[-*+]\s+', next_line):
                    i, nested_elements = process_list_items_returning_elements(
                        lines, i, doc, False, next_level, return_elements
                    )
                    if return_elements and nested_elements:
                        elements.extend(nested_elements)
                else:
                    # Not a list item, stop processing nested items
                    break
            elif next_level == level:
                # Same level item - continue in this loop
                break
            else:
                # Lower level - return to parent
                break

    return i, elements


# Regex patterns for block content detection (used by multiple modules)
ORDERED_LIST_PATTERN = re.compile(r'^\d+\.\s+')
UNORDERED_LIST_PATTERN = re.compile(r'^[-*+]\s+')
HEADING_PATTERN = re.compile(r'^(#{1,6})\s+(.+)$')


def contains_block_markdown(value: str) -> bool:
    """Check if the value contains block-level markdown content.

    Block-level content includes:
    - Ordered lists (1. item)
    - Unordered lists (- item, * item, + item)
    - Headings (# heading)

    Args:
        value: The string to check

    Returns:
        True if value contains block-level content
    """
    lines = value.split('\n')
    for line in lines:
        stripped = line.strip()
        if ORDERED_LIST_PATTERN.match(stripped):
            return True
        if UNORDERED_LIST_PATTERN.match(stripped):
            return True
        if HEADING_PATTERN.match(stripped):
            return True
    return False


def process_markdown_block(doc, lines, start_idx, return_element=True):
    """Process a single markdown block element (heading, list item start, or paragraph).

    Args:
        doc: The Word document
        lines: All lines of content
        start_idx: Current line index
        return_element: If True, remove element from doc and return it

    Returns:
        Tuple of (next_index, list_of_elements)
    """
    line = lines[start_idx]
    stripped = line.strip()
    elements = []

    # Check for heading
    heading_match = HEADING_PATTERN.match(stripped)
    if heading_match:
        level = len(heading_match.group(1))
        text = heading_match.group(2)
        heading = doc.add_heading('', level=min(level, 6))
        parse_inline_formatting(text, heading)
        if return_element:
            elements.append(heading._p)
            doc._body._body.remove(heading._p)
        return start_idx + 1, elements

    # Check for ordered list
    if ORDERED_LIST_PATTERN.match(stripped):
        return process_list_items_returning_elements(
            lines, start_idx, doc, is_ordered=True, level=0, return_elements=return_element
        )

    # Check for unordered list
    if UNORDERED_LIST_PATTERN.match(stripped):
        return process_list_items_returning_elements(
            lines, start_idx, doc, is_ordered=False, level=0, return_elements=return_element
        )

    # Regular paragraph
    para = doc.add_paragraph()
    parse_inline_formatting(stripped, para)
    if return_element:
        elements.append(para._p)
        doc._body._body.remove(para._p)
    return start_idx + 1, elements
